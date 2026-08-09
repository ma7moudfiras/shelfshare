# -*- coding: utf-8 -*-
"""Build the Aystro field training report as DOCX (English and Arabic)."""

import os
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

HERE = os.path.dirname(os.path.abspath(__file__))
FIGS = os.path.join(HERE, "figures")

INK = RGBColor(0x11, 0x11, 0x11)
MUTED = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0x1a, 0x3a, 0x5c)


# ----------------------------------------------------------------- low level

# OOXML enforces a strict child order inside these containers; inserting out of
# sequence produces a file Word tolerates but LibreOffice refuses to open.
RPR_SEQ = ("w:rStyle", "w:rFonts", "w:b", "w:bCs", "w:i", "w:iCs", "w:caps",
           "w:smallCaps", "w:strike", "w:dstrike", "w:outline", "w:shadow",
           "w:emboss", "w:imprint", "w:noProof", "w:snapToGrid", "w:vanish",
           "w:webHidden", "w:color", "w:spacing", "w:w", "w:kern", "w:position",
           "w:sz", "w:szCs", "w:highlight", "w:u", "w:effect", "w:bdr", "w:shd",
           "w:fitText", "w:vertAlign", "w:rtl", "w:cs", "w:em", "w:lang",
           "w:eastAsianLayout", "w:specVanish", "w:oMath")

PPR_SEQ = ("w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore",
           "w:framePr", "w:widowControl", "w:numPr", "w:suppressLineNumbers",
           "w:pBdr", "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku",
           "w:wordWrap", "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE",
           "w:autoSpaceDN", "w:bidi", "w:adjustRightInd", "w:snapToGrid",
           "w:spacing", "w:ind", "w:contextualSpacing", "w:mirrorIndents",
           "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment",
           "w:textboxTightWrap", "w:outlineLvl", "w:divId", "w:cnfStyle",
           "w:rPr", "w:sectPr", "w:pPrChange")

SECTPR_SEQ = ("w:footnotePr", "w:endnotePr", "w:type", "w:pgSz", "w:pgMar",
              "w:paperSrc", "w:pgBorders", "w:lnNumType", "w:pgNumType",
              "w:cols", "w:formProt", "w:vAlign", "w:noEndnote", "w:titlePg",
              "w:textDirection", "w:bidi", "w:rtlGutter", "w:docGrid",
              "w:printerSettings", "w:sectPrChange")

TBLPR_SEQ = ("w:tblStyle", "w:tblpPr", "w:tblOverlap", "w:bidiVisual",
             "w:tblStyleRowBandSize", "w:tblStyleColBandSize", "w:tblW",
             "w:jc", "w:tblCellSpacing", "w:tblInd", "w:tblBorders", "w:shd",
             "w:tblLayout", "w:tblCellMar", "w:tblLook", "w:tblCaption",
             "w:tblDescription")


def _ensure(parent, tag, seq):
    """Get or create `tag` inside `parent`, respecting the schema order."""
    el = parent.find(qn(tag))
    if el is not None:
        return el
    el = OxmlElement(tag)
    successors = seq[seq.index(tag) + 1:]
    parent.insert_element_before(el, *successors)
    return el


def _rtl_run(run):
    rPr = run._element.get_or_add_rPr()
    _ensure(rPr, "w:rtl", RPR_SEQ)
    _ensure(rPr, "w:cs", RPR_SEQ)


def _rtl_par(par):
    pPr = par._p.get_or_add_pPr()
    _ensure(pPr, "w:bidi", PPR_SEQ)
    for run in par.runs:
        _rtl_run(run)


def _set_fonts(run, latin, cs, size, bold=False, italic=False, color=INK):
    run.font.name = latin
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = _ensure(rPr, "w:rFonts", RPR_SEQ)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)
    rFonts.set(qn("w:cs"), cs)
    # complex-script size and bold must be set separately or Arabic ignores them
    _ensure(rPr, "w:szCs", RPR_SEQ).set(qn("w:val"), str(int(size * 2)))
    if bold:
        _ensure(rPr, "w:bCs", RPR_SEQ).set(qn("w:val"), "1")
    if italic:
        _ensure(rPr, "w:iCs", RPR_SEQ).set(qn("w:val"), "1")


class Builder:
    def __init__(self, rtl, body_font, cs_font, head_font):
        self.rtl = rtl
        self.body_font = body_font
        self.cs_font = cs_font
        self.head_font = head_font
        self.doc = Document()
        self._page_setup()
        self.fig_n = 0
        self.tab_n = 0

    def _page_setup(self):
        for s in self.doc.sections:
            s.page_width, s.page_height = Cm(21.0), Cm(29.7)
            s.left_margin = s.right_margin = Cm(2.6)
            s.top_margin = s.bottom_margin = Cm(2.4)
            if self.rtl:
                _ensure(s._sectPr, "w:bidi", SECTPR_SEQ)

    # ------------------------------------------------------------ primitives

    def para(self, text, size=11.5, bold=False, italic=False, color=INK,
             align=None, space_after=8, space_before=0, first_line=None,
             line=1.45, keep_with_next=False):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(space_after)
        pf.space_before = Pt(space_before)
        pf.line_spacing = line
        pf.keep_with_next = keep_with_next
        if align is not None:
            p.alignment = align
        elif self.rtl:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if first_line:
            pf.first_line_indent = Cm(first_line)
        r = p.add_run(text)
        _set_fonts(r, self.body_font, self.cs_font, size, bold, italic, color)
        if self.rtl:
            _rtl_par(p)
        return p

    def heading(self, text, level=1):
        size = 14 if level == 1 else 12
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(18 if level == 1 else 12)
        pf.space_after = Pt(7)
        pf.keep_with_next = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if self.rtl else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        _set_fonts(r, self.head_font, self.cs_font, size, True, False, ACCENT)
        if self.rtl:
            _rtl_par(p)
        return p

    def num(self, n):
        """Figure/table numbers must match the numeral system of the body text."""
        if not self.rtl:
            return str(n)
        return str(n).translate(str.maketrans("0123456789", "٠١٢٣٤٥٦٧٨٩"))

    def figure(self, name, caption, meta):
        path = next((os.path.join(FIGS, name + ext)
                     for ext in (".png", ".jpg")
                     if os.path.exists(os.path.join(FIGS, name + ext))), None)
        if path is None:
            raise FileNotFoundError(f"no figure image for {name!r} in {FIGS}")

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        # portrait figures must be capped by height or they overflow the page
        with PILImage.open(path) as probe:
            ratio = probe.height / probe.width
        if ratio > 1.15:
            p.add_run().add_picture(path, height=Cm(15.5))
        else:
            p.add_run().add_picture(path, width=Cm(15.4))

        self.fig_n += 1
        c = self.doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(14)
        lbl = c.add_run(f"{meta['fig']} {self.num(self.fig_n)}. ")
        _set_fonts(lbl, self.body_font, self.cs_font, 9.5, True, False, MUTED)
        txt = c.add_run(caption)
        # Arabic has no true italic; a synthesised oblique mangles the shaping
        _set_fonts(txt, self.body_font, self.cs_font, 9.5, False, not self.rtl, MUTED)
        if self.rtl:
            _rtl_par(c)

    def table(self, header, rows, caption, meta):
        self.tab_n += 1
        c = self.doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_before = Pt(12)
        c.paragraph_format.space_after = Pt(5)
        c.paragraph_format.keep_with_next = True
        lbl = c.add_run(f"{meta['table']} {self.num(self.tab_n)}. ")
        _set_fonts(lbl, self.body_font, self.cs_font, 9.5, True, False, MUTED)
        txt = c.add_run(caption)
        _set_fonts(txt, self.body_font, self.cs_font, 9.5, False, not self.rtl, MUTED)
        if self.rtl:
            _rtl_par(c)

        t = self.doc.add_table(rows=1, cols=len(header))
        t.style = "Table Grid"
        t.alignment = 2 if self.rtl else 0
        if self.rtl:
            _ensure(t._tbl.tblPr, "w:bidiVisual", TBLPR_SEQ)

        def cell_text(cell, text, bold):
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if self.rtl else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(text)
            _set_fonts(r, self.body_font, self.cs_font, 10, bold)
            if self.rtl:
                _rtl_par(p)

        for i, h in enumerate(header):
            cell_text(t.rows[0].cells[i], h, True)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cell_text(cells[i], v, False)

        sp = self.doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(10)
        sp.paragraph_format.space_before = Pt(0)

    def page_break(self):
        p = self.doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

    # ----------------------------------------------------------- title page

    def title_page(self, meta):
        for _ in range(3):
            self.doc.add_paragraph()
        self.para(meta["org"], size=13, bold=True, color=MUTED,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
        self.para(meta["title"], size=21, bold=True, color=ACCENT,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, line=1.25)
        self.para(meta["subtitle"], size=14, italic=not self.rtl, color=MUTED,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=64)

        L = meta["labels"]
        for label, value in ((L["trainee"], meta["author"]),
                             (L["org"], meta["org"]),
                             (L["period"], meta["period"]),
                             (L["submitted"], meta["submitted"])):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(7)
            a = p.add_run(f"{label}:  ")
            _set_fonts(a, self.body_font, self.cs_font, 11.5, True, False, MUTED)
            b = p.add_run(value)
            _set_fonts(b, self.body_font, self.cs_font, 11.5, False, False, INK)
            if self.rtl:
                _rtl_par(p)
        self.page_break()

    # ---------------------------------------------------------------- build

    def build(self, meta, abstract, sections, references):
        self.title_page(meta)

        self.heading(meta["abstract_h"], 1)
        for t in abstract:
            self.para(t, size=11)
        self.page_break()

        for title, blocks in sections:
            self.heading(title, 1)
            for block in blocks:
                kind = block[0]
                if kind == "p":
                    self.para(block[1])
                elif kind == "b":
                    self.para(block[1], space_after=6)
                elif kind == "h2":
                    self.heading(block[1], 2)
                elif kind == "fig":
                    self.figure(block[1], block[2], meta)
                elif kind == "table":
                    self.table(block[1], block[2], block[3], meta)

        self.page_break()
        self.heading(meta["refs_h"], 1)
        for i, ref in enumerate(references, 1):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.space_after = Pt(8)
            pf.line_spacing = 1.3
            pf.left_indent = Cm(0.9)
            pf.first_line_indent = Cm(-0.9)
            r = p.add_run(f"[{i}]  {ref}")
            _set_fonts(r, self.body_font, "Times New Roman", 10.5)
        return self.doc


def main():
    import content_en, content_ar

    jobs = [
        (content_en, False, "Times New Roman", "Times New Roman",
         "Times New Roman", "Aystro_Field_Training_Report_EN"),
        (content_ar, True, "Times New Roman", "Noto Naskh Arabic",
         "Times New Roman", "Aystro_Field_Training_Report_AR"),
    ]
    out = os.path.join(HERE, "out")
    os.makedirs(out, exist_ok=True)

    for mod, rtl, body, cs, head, name in jobs:
        b = Builder(rtl, body, cs, head)
        doc = b.build(mod.META, mod.ABSTRACT, mod.SECTIONS, mod.REFERENCES)
        path = os.path.join(out, name + ".docx")
        doc.save(path)
        print("wrote", path)


if __name__ == "__main__":
    main()
